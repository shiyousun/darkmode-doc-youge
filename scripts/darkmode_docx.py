#!/usr/bin/env python3
"""
darkmode_docx.py — 把 Word (.docx) 转成护眼模式

两档模式（v2.0 新增）：
- 默认（不加任何参数）：只改文档背景 + 文字 + 表格 + 边框，**不改插图**
- 加 `--invert-images`：同时温柔反色所有内嵌图片（含 PNG/JPEG/BMP/GIF/TIFF/WEBP，
  以及 WMF 矢量公式图，需 libwmf 的 wmf2gd 工具）

处理项：
- 页面背景改主题深色（启用 displayBackgroundShape）
- 所有正文 / 表格 / 页眉 / 页脚 中的文字改主题前景色
- 表格单元格底色改主题深色
- 表格边框改主题边框色
- (仅 invert_images=True) 内嵌图片温柔反色 + WMF 用 wmf2gd 转 PNG 后反色

用法：
    python3 darkmode_docx.py input.docx [output.docx] [--invert-images] [--theme warm]

注意：
- `--no-images` 保留作为向后兼容（现在与默认行为一致）
- 显式启用 `--invert-images` 才反色图片
"""
from __future__ import annotations

import argparse
import io
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

SCRIPT_DIR = Path(__file__).resolve().parent
sys.path.insert(0, str(SCRIPT_DIR))

from themes import DEFAULT_THEME, THEMES, gentle_invert, get_theme


WMF2GD_CANDIDATES = [
    "/opt/homebrew/opt/libwmf/bin/wmf2gd",
    "/usr/local/opt/libwmf/bin/wmf2gd",
    "/opt/homebrew/bin/wmf2gd",
    "/usr/local/bin/wmf2gd",
    "/usr/bin/wmf2gd",
]


def find_wmf2gd() -> str | None:
    for cand in WMF2GD_CANDIDATES:
        if Path(cand).exists():
            return cand
    found = shutil.which("wmf2gd")
    return found


def _ensure_setting(settings_element, tag: str) -> None:
    el = settings_element.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        settings_element.append(el)


def set_page_background(doc, color_hex: str) -> None:
    """设置 Word 页面背景色 + 启用打开时显示背景。"""
    settings = doc.settings.element
    _ensure_setting(settings, "w:displayBackgroundShape")

    document_element = doc.element
    background = document_element.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        document_element.insert(0, background)
    background.set(qn("w:color"), color_hex)


def set_run_color(run, color_hex: str) -> None:
    """把一个 run 的字体颜色设置为指定色。"""
    rPr = run._element.get_or_add_rPr()
    color_el = rPr.find(qn("w:color"))
    if color_el is None:
        color_el = OxmlElement("w:color")
        rPr.append(color_el)
    color_el.set(qn("w:val"), color_hex)
    if color_el.get(qn("w:themeColor")) is not None:
        del color_el.attrib[qn("w:themeColor")]


def set_paragraph_default_color(paragraph, color_hex: str) -> None:
    """段落默认字体色（兜底没有 run 的情况）。"""
    pPr = paragraph._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        pPr.append(rPr)
    color_el = rPr.find(qn("w:color"))
    if color_el is None:
        color_el = OxmlElement("w:color")
        rPr.append(color_el)
    color_el.set(qn("w:val"), color_hex)


def set_cell_background(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)


def set_table_borders(table, color_hex: str) -> None:
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for tag in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{tag}"))
        if el is None:
            el = OxmlElement(f"w:{tag}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)


def process_paragraph(paragraph, font_color: str) -> None:
    set_paragraph_default_color(paragraph, font_color)
    for run in paragraph.runs:
        set_run_color(run, font_color)


def process_table(table, font_color: str, cell_bg: str, border_color: str) -> None:
    set_table_borders(table, border_color)
    for row in table.rows:
        for cell in row.cells:
            if cell_bg:
                set_cell_background(cell, cell_bg)
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph, font_color)
            for nested in cell.tables:
                process_table(nested, font_color, cell_bg, border_color)


def invert_image_blob(blob: bytes, fmt_hint: str, theme: dict) -> bytes | None:
    """对 docx 内嵌图片做温柔反色。WMF/EMF 在此处不处理（由 _post_process_wmf 接管）。
    失败返回 None。"""
    fmt = fmt_hint.lower().lstrip(".")
    if fmt in ("wmf", "emf", "x-wmf", "x-emf"):
        return None
    try:
        img = Image.open(io.BytesIO(blob))
        warm = gentle_invert(img, theme)

        save_format = {
            "jpg": "JPEG",
            "jpeg": "JPEG",
            "png": "PNG",
            "bmp": "BMP",
            "gif": "GIF",
            "tif": "TIFF",
            "tiff": "TIFF",
            "webp": "WEBP",
        }.get(fmt)

        if save_format is None:
            save_format = (img.format or "PNG").upper()

        if save_format == "JPEG" and warm.mode in ("RGBA", "LA"):
            warm = warm.convert("RGB")

        out = io.BytesIO()
        save_kwargs = {}
        if save_format == "JPEG":
            save_kwargs["quality"] = 90
        warm.save(out, format=save_format, **save_kwargs)
        return out.getvalue()
    except Exception as e:
        print(f"  ! 图片反色失败 ({fmt_hint}): {e}", flush=True)
        return None


def invert_docx_images(doc, theme: dict) -> int:
    """遍历 docx 中所有 image part，做温柔反色。返回成功反色的图片数。
    跳过 WMF/EMF（由 _post_process_wmf 接管）。"""
    seen = set()
    counter = {"n": 0}

    def visit(part):
        if id(part) in seen:
            return
        seen.add(id(part))
        for rel in list(part.rels.values()):
            if rel.is_external:
                continue
            try:
                target = rel.target_part
            except Exception:
                continue
            content_type = getattr(target, "content_type", "") or ""
            if content_type.startswith("image/"):
                ext = content_type.split("/")[-1]
                blob = target.blob
                new_blob = invert_image_blob(blob, ext, theme)
                if new_blob:
                    target._blob = new_blob
                    counter["n"] += 1
            else:
                visit(target)

    visit(doc.part)
    return counter["n"]


def _post_process_wmf(docx_path: Path, theme: dict, wmf2gd: str) -> int:
    """
    对已生成的 _dark.docx 做 WMF 反色后处理。
    解压 → wmf2gd 把 WMF 转 PNG → 温柔反色 → 改 [Content_Types].xml 与 *.rels
    → 删除原 WMF → 重新打包。

    返回成功反色的 WMF 数。0 表示无 WMF 或全部失败。
    """
    tmpdir = Path(tempfile.mkdtemp(prefix="docx_wmf_"))
    try:
        extract_dir = tmpdir / "extract"
        extract_dir.mkdir()
        with zipfile.ZipFile(docx_path, "r") as z:
            z.extractall(extract_dir)

        wmf_files = sorted(extract_dir.rglob("*.wmf")) + sorted(extract_dir.rglob("*.WMF"))
        if not wmf_files:
            return 0

        media_dir = wmf_files[0].parent
        png_out = tmpdir / "png_out"
        png_out.mkdir()

        for wmf in wmf_files:
            out_png = png_out / (wmf.stem + ".png")
            subprocess.run(
                [
                    wmf2gd,
                    "-t",
                    "png",
                    "--maxwidth=2400",
                    "--maxheight=2400",
                    "-o",
                    str(out_png),
                    str(wmf),
                ],
                capture_output=True,
                timeout=30,
            )

        wmf_to_png: dict[str, str] = {}
        inverted_count = 0
        for wmf in wmf_files:
            png_path = png_out / (wmf.stem + ".png")
            if not png_path.exists() or png_path.stat().st_size == 0:
                continue
            try:
                with Image.open(png_path) as img:
                    if img.mode in ("RGBA", "LA", "P"):
                        bg = Image.new("RGB", img.size, (255, 255, 255))
                        if img.mode == "P":
                            img = img.convert("RGBA")
                        bg.paste(
                            img,
                            mask=img.split()[-1] if img.mode in ("RGBA", "LA") else None,
                        )
                        img = bg
                    elif img.mode != "RGB":
                        img = img.convert("RGB")
                    warm = gentle_invert(img, theme)
                    new_png = media_dir / (wmf.stem + ".png")
                    warm.save(new_png, format="PNG", optimize=True)
                wmf_to_png[wmf.name] = wmf.stem + ".png"
                inverted_count += 1
            except Exception as e:
                print(f"  ! WMF 反色失败 {wmf.name}: {e}", flush=True)

        if not wmf_to_png:
            return 0

        ct_path = extract_dir / "[Content_Types].xml"
        if ct_path.exists():
            ct = ct_path.read_text(encoding="utf-8")
            if 'Extension="png"' not in ct:
                ct = ct.replace(
                    "</Types>",
                    '<Default Extension="png" ContentType="image/png"/></Types>',
                )
                ct_path.write_text(ct, encoding="utf-8")

        for rels_path in extract_dir.rglob("*.rels"):
            text = rels_path.read_text(encoding="utf-8")
            new_text = text
            for wmf_name, png_name in wmf_to_png.items():
                new_text = new_text.replace(wmf_name, png_name)
                new_text = new_text.replace(
                    wmf_name.replace(".wmf", ".WMF"), png_name
                )
            if new_text != text:
                rels_path.write_text(new_text, encoding="utf-8")

        for wmf in wmf_files:
            try:
                wmf.unlink()
            except Exception:
                pass

        new_docx = tmpdir / "rebuilt.docx"
        with zipfile.ZipFile(new_docx, "w", zipfile.ZIP_DEFLATED, compresslevel=6) as zout:
            for f in sorted(extract_dir.rglob("*")):
                if f.is_file():
                    arcname = f.relative_to(extract_dir).as_posix()
                    zout.write(f, arcname)
        shutil.move(str(new_docx), str(docx_path))
        return inverted_count
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def process_section_part(
    part_obj, font_color: str, cell_bg: str, border_color: str
) -> None:
    """处理 header/footer。"""
    if part_obj is None:
        return
    for paragraph in getattr(part_obj, "paragraphs", []):
        process_paragraph(paragraph, font_color)
    for table in getattr(part_obj, "tables", []):
        process_table(table, font_color, cell_bg, border_color)


def convert_docx_to_darkmode(
    input_path: Path,
    output_path: Path,
    invert_images: bool = False,
    theme_name: str = DEFAULT_THEME,
) -> dict:
    """
    把 .docx 转成护眼模式。

    参数：
        invert_images:
            False（默认）— 只改背景/文字/表格，不动插图
            True         — 同时温柔反色所有内嵌图片（含 WMF，需 wmf2gd）
    """
    theme = get_theme(theme_name)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(str(input_path), str(output_path))

    doc = Document(str(output_path))

    bg = theme["bg"]
    fg = theme["fg"]
    border = theme["border"]

    set_page_background(doc, bg)

    for paragraph in doc.paragraphs:
        process_paragraph(paragraph, fg)

    for table in doc.tables:
        process_table(table, fg, bg, border)

    for section in doc.sections:
        for hf_attr in (
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ):
            try:
                hf = getattr(section, hf_attr)
            except Exception:
                hf = None
            process_section_part(hf, fg, bg, border)

    images_inverted = 0
    wmf_inverted = 0
    if invert_images:
        images_inverted = invert_docx_images(doc, theme)

    doc.save(str(output_path))

    if invert_images:
        wmf2gd = find_wmf2gd()
        if wmf2gd:
            wmf_inverted = _post_process_wmf(output_path, theme, wmf2gd)

    return {
        "images_inverted": images_inverted,
        "wmf_inverted": wmf_inverted,
        "theme": theme_name,
        "size_in": input_path.stat().st_size,
        "size_out": output_path.stat().st_size,
    }


def parse_args(argv: list[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="DOCX 护眼模式转换器")
    parser.add_argument("input", help="输入 .docx 路径")
    parser.add_argument("output", nargs="?", help="输出路径（默认 xxx_dark.docx）")
    parser.add_argument(
        "--invert-images",
        action="store_true",
        help="同时温柔反色内嵌图片（含 WMF 公式）。默认不反色插图。",
    )
    parser.add_argument(
        "--no-images",
        action="store_true",
        help="（兼容旧版，与默认行为一致：不反色插图）",
    )
    parser.add_argument(
        "--theme",
        default=DEFAULT_THEME,
        choices=list(THEMES.keys()),
        help=f"护眼主题（默认 {DEFAULT_THEME}）",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv if argv is not None else sys.argv[1:])
    input_path = Path(args.input).expanduser().resolve()
    if not input_path.exists():
        print(f"❌ 输入文件不存在: {input_path}", file=sys.stderr)
        return 2

    if args.output:
        output_path = Path(args.output).expanduser().resolve()
    else:
        output_path = input_path.with_name(f"{input_path.stem}_dark.docx")

    invert_images = bool(args.invert_images) and not bool(args.no_images)

    theme = get_theme(args.theme)
    print(f"▶ DOCX 护眼模式转换")
    print(f"  输入: {input_path}")
    print(f"  输出: {output_path}")
    print(f"  反色内嵌图片: {'是' if invert_images else '否（仅改背景+文字）'}")
    if invert_images:
        wmf2gd = find_wmf2gd()
        print(f"  WMF 公式反色: {'是 (' + wmf2gd + ')' if wmf2gd else '否（未找到 wmf2gd，brew install libwmf 启用）'}")
    print(f"  主题: {args.theme} ({theme['name']}) bg=#{theme['bg']} fg=#{theme['fg']}")

    info = convert_docx_to_darkmode(
        input_path,
        output_path,
        invert_images=invert_images,
        theme_name=args.theme,
    )
    extras = []
    if invert_images:
        extras.append(f"反色图片 {info['images_inverted']} 张")
        if info.get("wmf_inverted"):
            extras.append(f"反色 WMF {info['wmf_inverted']} 张")
    extras.append(
        f"{info['size_in'] / 1024 / 1024:.2f} MB → {info['size_out'] / 1024 / 1024:.2f} MB"
    )
    print(f"✅ 完成 · " + " · ".join(extras))
    return 0


if __name__ == "__main__":
    sys.exit(main())
